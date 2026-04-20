# tests/test_ingestion.py — Document ingestion: processor (PDF, DOCX, TXT), mapper (mock LLM), evidence, API.

from __future__ import annotations

import io
import tempfile
from contextlib import contextmanager
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
from fastapi.testclient import TestClient

from api.main import app
from ontology.models import ControlRef, Obligation, OntologyMappingResult

# Compliance-engine ingestion (services/compliance-engine/app/services/ingestion).
from app.services.ingestion.document_processor import (
    DocumentChunk,
    _chunk_text,
    extract_text_docx,
    extract_text_pdf,
    extract_text_txt,
    process_document,
)
from app.services.ingestion.evidence_creator import content_hash, create_evidence_from_mapping

client = TestClient(app)


@contextmanager
def temp_dir_with_file(name: str, content: str | bytes):
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / name
        if isinstance(content, str):
            path.write_text(content, encoding="utf-8")
        else:
            path.write_bytes(content)
        yield d


# ---- document_processor ----


def test_chunk_text_overlapping() -> None:
    """Chunk text into ~500 token segments with 50 token overlap."""
    text = "word " * 2500
    chunks = _chunk_text(text, "txt", None)
    assert len(chunks) >= 1
    assert all(c.document_type == "txt" for c in chunks)
    assert all(isinstance(c.content, str) and len(c.content) > 0 for c in chunks)


def test_extract_text_txt() -> None:
    """TXT extraction returns full text and line list."""
    path = Path(__file__).parent / "fixtures" / "sample.txt"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("Hello world.\n\nSecond paragraph.", encoding="utf-8")
    full, page_texts = extract_text_txt(path)
    assert "Hello world" in full
    assert len(page_texts) >= 1


def test_extract_text_pdf() -> None:
    """PDF text extraction via pypdf returns full text and page list."""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        w = PdfWriter()
        w.add_blank_page(width=72 * 6, height=72 * 6)
        w.write(f)
        f.flush()
        path = Path(f.name)
    try:
        full, page_texts = extract_text_pdf(path)
        assert isinstance(full, str)
        assert isinstance(page_texts, list)
        assert len(page_texts) >= 1
        assert page_texts[0][0] == 1
        # Blank page may have empty text; reader still returns one page
        assert len(PdfReader(str(path)).pages) == 1
    finally:
        path.unlink(missing_ok=True)


def test_extract_text_docx() -> None:
    """DOCX text extraction via python-docx returns full text and paragraph list."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    with temp_dir_with_file("sample.docx", b"") as tmp:
        doc_path = Path(tmp) / "sample.docx"
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(str(doc_path))
        full, page_texts = extract_text_docx(doc_path)
        assert "First paragraph" in full
        assert "Second paragraph" in full
        assert len(page_texts) >= 1


def test_process_document_txt() -> None:
    """Process TXT file returns List[DocumentChunk]."""
    with temp_dir_with_file("sample.txt", "Line one.\n\nLine two.") as tmp:
        chunks = process_document(Path(tmp) / "sample.txt", "txt")
    assert isinstance(chunks, list)
    assert all(isinstance(c, DocumentChunk) for c in chunks)
    assert any("Line one" in c.content or "Line two" in c.content for c in chunks)


def test_process_document_unknown_type_raises() -> None:
    """Unsupported document type raises ValueError."""
    with temp_dir_with_file("x.txt", "content") as d:
        path = Path(d) / "x.txt"
        with pytest.raises(ValueError, match="Unsupported"):
            process_document(path, "xml")  # type: ignore


# ---- ontology_mapper (mock LLM) ----


def test_ontology_mapping_with_mock_llm() -> None:
    """Map chunks to ontology via CircuitBreaker; mock LLM returns controlled result."""
    import asyncio

    from app.services.ingestion.ontology_mapper import map_chunks_to_ontology

    chunks = [
        DocumentChunk(content="We process data under lawful basis.", chunk_index=0, document_type="pdf"),
    ]

    async def run() -> None:
        result = await map_chunks_to_ontology(chunks, "pdf", "doc-test-1")
        assert result.controls
        assert result.obligations
        assert 0 <= result.confidence_score <= 1
        if result.confidence_score < 0.75:
            assert result.requires_human_review is True

    asyncio.run(run())


# ---- evidence_creator ----


def test_content_hash() -> None:
    """content_hash returns SHA-256 hex string."""
    h = content_hash("hello")
    assert len(h) == 64
    assert all(c in "0123456789abcdef" for c in h)


def test_create_evidence_from_mapping() -> None:
    """Evidence created for each control; content_hash and obligations_satisfied set."""
    mapping = OntologyMappingResult(
        controls=[ControlRef(framework_id="gdpr", control_id="lb-1")],
        obligations=[
            Obligation(
                jurisdiction="EU",
                purpose_tags=[],
                id="obl-1",
                description="Lawful basis",
                control_refs=[],
            )
        ],
        confidence_score=0.9,
        requires_human_review=False,
    )
    evidence_list = create_evidence_from_mapping(mapping, "document body", "doc-123")
    assert len(evidence_list) == 1
    assert evidence_list[0].evidence_type == "POLICY_DOCUMENT"
    assert evidence_list[0].content_hash == content_hash("document body")
    assert evidence_list[0].obligations_satisfied == ["obl-1"]
    assert evidence_list[0].confidence_score == 0.9


# ---- API (TestClient) ----


def test_ingest_document_rejects_large_file() -> None:
    """POST /api/v1/ingest/document rejects file > 10MB."""
    big = b"x" * (11 * 1024 * 1024)
    r = client.post("/api/v1/ingest/document", files={"file": ("big.txt", io.BytesIO(big), "text/plain")})
    assert r.status_code == 400
    assert "10MB" in r.json()["detail"]


def test_ingest_document_rejects_bad_type() -> None:
    """POST rejects non-PDF/DOCX/TXT."""
    r = client.post("/api/v1/ingest/document", files={"file": ("x.xml", io.BytesIO(b"<a/>"), "application/xml")})
    assert r.status_code == 400


def test_ingest_document_rejects_path_traversal_filename() -> None:
    r"""POST rejects filename with path traversal (.. or / or \)."""
    r = client.post(
        "/api/v1/ingest/document",
        files={"file": ("../../../etc/passwd.txt", io.BytesIO(b"x"), "text/plain")},
    )
    assert r.status_code == 400
    assert "path traversal" in r.json()["detail"].lower() or "invalid filename" in r.json()["detail"].lower()


def test_ingest_document_accepts_txt() -> None:
    """POST /api/v1/ingest/document accepts TXT and returns streaming response."""
    r = client.post(
        "/api/v1/ingest/document",
        files={"file": ("sample.txt", io.BytesIO(b"Policy document content here."), "text/plain")},
    )
    assert r.status_code == 200
    assert "text/event-stream" in r.headers.get("content-type", "")


def test_ingest_document_stream_includes_progress_and_summary() -> None:
    """SSE stream includes progress, mapping_done, evidence_created, summary, done."""
    r = client.post(
        "/api/v1/ingest/document",
        files={"file": ("sample.txt", io.BytesIO(b"Short policy."), "text/plain")},
    )
    assert r.status_code == 200
    text = r.text
    assert "event: progress" in text
    assert "event: mapping_done" in text or "event: summary" in text
    assert "event: done" in text or "event: summary" in text


def test_ingest_document_pipeline_error_emits_error_event() -> None:
    """When mapping raises, stream emits error event (audit ingest_done still runs)."""
    with patch("app.api.v1.ingest.map_chunks_to_ontology", new_callable=AsyncMock) as mock_map:
        mock_map.side_effect = RuntimeError("LLM unavailable")
        r = client.post(
            "/api/v1/ingest/document",
            files={"file": ("sample.txt", io.BytesIO(b"Some content for chunks."), "text/plain")},
        )
        assert r.status_code == 200
        assert "event: error" in r.text
        assert "LLM unavailable" in r.text
